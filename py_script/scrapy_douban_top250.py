#!/usr/bin/env python
# coding=utf-8
import requests
from bs4 import BeautifulSoup
from docx import Document
import os

doc = Document()
doc.add_heading("豆瓣电影Top250",0)
url = 'https://movie.douban.com/top250?start=0&filter='
web_data = requests.get(url);
soup = BeautifulSoup(web_data.text, 'lxml')
imgs = soup.select('img');
for img in imgs:
    print(img.get('alt'),img.get('src'))
    contents = str(img.get('alt')) + str(img.get('src'))
    doc.add_paragraph(contents)

doc.save("E:\\DannyWu\\Desktop\\doubanTop250.docx")

os.system("start explorer E:\\DannyWu\\Desktop")
os.system("start https://movie.douban.com/top250?start=0&filter=")
#print(soup.prettify())
'''



肖申克的救赎 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p480747492.jpg
霸王别姬 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1910813120.jpg
这个杀手不太冷 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p511118051.jpg
阿甘正传 https://img1.doubanio.com/view/photo/s_ratio_poster/public/p510876377.jpg
美丽人生 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p510861873.jpg
千与千寻 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1606727862.jpg
泰坦尼克号 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p457760035.jpg
辛德勒的名单 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p492406163.jpg
盗梦空间 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p513344864.jpg
机器人总动员 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1461851991.jpg
三傻大闹宝莱坞 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p579729551.jpg
海上钢琴师 https://img1.doubanio.com/view/photo/s_ratio_poster/public/p511146807.jpg
忠犬八公的故事 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p524964016.jpg
放牛班的春天 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1910824951.jpg
大话西游之大圣娶亲 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p2455050536.jpg
楚门的世界 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p479682972.jpg
龙猫 https://img1.doubanio.com/view/photo/s_ratio_poster/public/p537668599.jpg
教父 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p2190556185.jpg
星际穿越 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p2206088801.jpg
熔炉 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1363250216.jpg
触不可及 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1454261925.jpg
乱世佳人 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1963126880.jpg
无间道 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p2233971046.jpg
当幸福来敲门 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p1312700744.jpg
怦然心动 https://img3.doubanio.com/view/photo/s_ratio_poster/public/p663036666.jpg

'''

